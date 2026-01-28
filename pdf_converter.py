import os
import pdfplumber
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io
import re


def _clean_text(text):
    """
    Clean text to remove XML-incompatible characters.
    Removes NULL bytes, control characters, and other problematic characters.
    """
    if not text:
        return ""

    # Remove NULL bytes
    text = text.replace('\x00', '')

    # Remove other control characters except newline, tab, and carriage return
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')

    # Remove any remaining problematic Unicode characters
    text = text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')

    return text.strip()


def _set_font(run, is_code=False):
    """
    Set appropriate fonts for text to ensure compatibility with Microsoft Word.
    """
    if is_code:
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    else:
        run.font.name = 'Calibri'
        run.font.size = Pt(11)  # Set default font size
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')


def convert_pdf_to_docx(pdf_file_path, output_path, extract_images=True):
    """
    Convert a PDF file to a Word document with improved formatting.

    Args:
        pdf_file_path: Path to the PDF file
        output_path: Path where the Word document should be saved
        extract_images: Whether to extract and embed images from PDF
    """
    # Create Word document
    doc = Document()

    # Open PDF file
    with pdfplumber.open(pdf_file_path) as pdf:
        print(f"[DEBUG] Processing PDF with {len(pdf.pages)} pages")

        for page_num, page in enumerate(pdf.pages, 1):
            print(f"[DEBUG] Processing page {page_num}/{len(pdf.pages)}")

            # Add page break (except for first page)
            if page_num > 1:
                doc.add_page_break()

            # Get page layout info
            page_width = page.width
            page_height = page.height

            # Extract images first if requested
            images_positions = []
            if extract_images:
                try:
                    images = page.images
                    if images:
                        print(f"[DEBUG] Found {len(images)} images on page {page_num}")
                        for img in images:
                            images_positions.append({
                                'x0': img['x0'],
                                'top': img['top'],
                                'x1': img['x1'],
                                'bottom': img['bottom']
                            })
                except Exception as e:
                    print(f"[DEBUG] Error getting images info: {e}")

            # Extract text with layout
            text = page.extract_text(layout=True)
            if text:
                # Clean text
                text = _clean_text(text)

                # Process text line by line to preserve layout
                lines = text.split('\n')
                current_para = []

                for line in lines:
                    line = _clean_text(line)
                    if not line:
                        # Empty line - end current paragraph
                        if current_para:
                            para_text = ' '.join(current_para)
                            _add_paragraph_with_style(doc, para_text)
                            current_para = []
                        continue

                    # Check if line is a heading
                    if _is_heading(line):
                        # End current paragraph first
                        if current_para:
                            para_text = ' '.join(current_para)
                            _add_paragraph_with_style(doc, para_text)
                            current_para = []

                        # Add heading
                        heading = doc.add_heading(line, level=_get_heading_level(line))
                        for run in heading.runs:
                            _set_font(run)
                    else:
                        # Regular text - accumulate into paragraph
                        current_para.append(line)

                # Add remaining paragraph
                if current_para:
                    para_text = ' '.join(current_para)
                    _add_paragraph_with_style(doc, para_text)

            # Extract and add tables
            tables = page.extract_tables()
            if tables:
                print(f"[DEBUG] Found {len(tables)} tables on page {page_num}")
                for table_data in tables:
                    _add_table_to_doc(doc, table_data)

            # Extract and add images
            if extract_images and images_positions:
                for img_index, img_pos in enumerate(images_positions):
                    try:
                        _extract_and_add_image(doc, page, img_pos, page_num, img_index)
                    except Exception as e:
                        print(f"[DEBUG] Error extracting image {img_index}: {e}")

    # Save document
    doc.save(output_path)
    print(f"[DEBUG] PDF converted successfully to {output_path}")


def _add_paragraph_with_style(doc, text):
    """Add a paragraph with proper styling."""
    if not text or not text.strip():
        return

    p = doc.add_paragraph(text)
    for run in p.runs:
        _set_font(run)

    # Set paragraph spacing
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15


def _is_heading(text):
    """
    Detect if text looks like a heading with improved logic.
    """
    text = text.strip()
    if not text or len(text) > 150:
        return False

    # Check for numbered headings (1. 1.1 第一章 etc.)
    if re.match(r'^[\d一二三四五六七八九十]+[\.、\s]', text):
        return True

    # Check for chapter/section markers
    if re.match(r'^(第[一二三四五六七八九十\d]+[章节部分]|Chapter|Section|CHAPTER|SECTION)', text):
        return True

    # Check if mostly uppercase and short
    if len(text) < 50 and sum(1 for c in text if c.isupper()) / len(text) > 0.7:
        return True

    # Check for common heading patterns
    if text.endswith('：') or text.endswith(':'):
        return True

    return False


def _get_heading_level(text):
    """Determine heading level based on text content."""
    text = text.strip()

    # Level 1: Chapter markers
    if re.match(r'^(第[一二三四五六七八九十\d]+章|Chapter\s+\d+|CHAPTER\s+\d+)', text):
        return 1

    # Level 2: Section markers or numbered like 1. 2.
    if re.match(r'^[\d一二三四五六七八九十]+[\.、]\s', text):
        # Count dots to determine level
        dots = text.count('.')
        if dots == 0:
            return 2
        elif dots == 1:
            return 3
        else:
            return 4

    # Level 2: Section markers
    if re.match(r'^(第[一二三四五六七八九十\d]+节|Section)', text):
        return 2

    # Default to level 3
    return 3


def _add_table_to_doc(doc, table_data):
    """
    Add a table to the Word document with improved formatting.
    """
    if not table_data or len(table_data) == 0:
        return

    # Filter out None rows and get dimensions
    valid_rows = [row for row in table_data if row and any(cell for cell in row)]
    if not valid_rows:
        return

    num_rows = len(valid_rows)
    num_cols = max(len(row) for row in valid_rows)

    # Create Word table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Fill table
    for i, row in enumerate(valid_rows):
        for j, cell in enumerate(row):
            if j < num_cols and cell is not None:
                # Clean cell text
                cell_text = _clean_text(str(cell))
                table.rows[i].cells[j].text = cell_text

                # Set font for table cells
                for paragraph in table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        _set_font(run)

                # Make first row bold (header)
                if i == 0:
                    for paragraph in table.rows[i].cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    # Add spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)


def _extract_and_add_image(doc, page, img_pos, page_num, img_index):
    """
    Extract image from PDF page and add to Word document with improved handling.
    """
    try:
        x0, y0, x1, y1 = img_pos['x0'], img_pos['top'], img_pos['x1'], img_pos['bottom']

        # Skip very small images (likely decorative)
        width = x1 - x0
        height = y1 - y0
        if width < 20 or height < 20:
            print(f"[DEBUG] Skipping small image {img_index} ({width}x{height})")
            return

        # Crop image from page
        bbox = (x0, y0, x1, y1)
        cropped_page = page.crop(bbox)

        # Convert to image with higher resolution
        img = cropped_page.to_image(resolution=200)

        # Save to bytes
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)

        # Calculate width (max 6 inches, maintain aspect ratio)
        img_width_inches = min(6, width / 72)  # Convert points to inches

        # Add to document
        doc.add_picture(img_bytes, width=Inches(img_width_inches))

        # Add spacing after image
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

        print(f"[DEBUG] Added image {img_index} from page {page_num} ({width:.0f}x{height:.0f})")

    except Exception as e:
        print(f"[DEBUG] Could not extract image {img_index} from page {page_num}: {e}")


def convert_pdf_to_docx_simple(pdf_file_path, output_path):
    """
    Simple PDF to Word conversion with improved formatting.
    """
    # Use the full conversion with images enabled
    convert_pdf_to_docx(pdf_file_path, output_path, extract_images=True)

