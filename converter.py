import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import markdown2
from bs4 import BeautifulSoup
from pathlib import Path
from urllib.parse import unquote


def _set_font(run, is_code=False):
    """
    Set appropriate fonts for text to ensure compatibility with Microsoft Word.

    Args:
        run: The run object to set font for
        is_code: Whether this is code text (uses monospace font)
    """
    if is_code:
        # Code blocks use monospace font
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        # Set East Asian font for Chinese characters in code
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    else:
        # Regular text uses standard fonts
        run.font.name = 'Calibri'
        # Set East Asian font for Chinese/Japanese/Korean characters
        # This ensures proper display in Microsoft Word
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        # Set fallback fonts for better compatibility
        run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
        run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')


def convert_markdown_to_docx(md_file_path, output_path, images_dir=None):
    """
    Convert a markdown file to a Word document with advanced formatting.

    Args:
        md_file_path: Path to the markdown file
        output_path: Path where the Word document should be saved
        images_dir: Directory containing images referenced in markdown
    """
    # Read markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Convert markdown to HTML with extras for tables, code blocks, etc.
    html = markdown2.markdown(
        md_content,
        extras=[
            'tables',
            'fenced-code-blocks',
            'code-friendly',
            'break-on-newline',
            'task_list'
        ]
    )

    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(html, 'html5lib')

    # Create Word document
    doc = Document()

    # Process each element in the HTML
    _process_element(doc, soup.body, images_dir)

    # Save document
    doc.save(output_path)


def _process_element(doc, element, images_dir, list_level=0):
    """
    Recursively process HTML elements and convert to Word document elements.
    """
    if element is None:
        return

    for child in element.children:
        if isinstance(child, str):
            # Skip whitespace-only text nodes at document level
            if child.strip():
                p = doc.add_paragraph(child.strip())
                # Set font for all runs in the paragraph
                for run in p.runs:
                    _set_font(run)
            continue

        tag_name = child.name

        # Headings
        if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag_name[1])
            text = child.get_text()
            heading = doc.add_heading(text, level=level)
            # Set font for heading
            for run in heading.runs:
                _set_font(run)

        # Paragraphs
        elif tag_name == 'p':
            p = doc.add_paragraph()
            _process_inline_elements(p, child, images_dir)

        # Unordered lists
        elif tag_name == 'ul':
            _process_list(doc, child, images_dir, ordered=False, level=list_level)

        # Ordered lists
        elif tag_name == 'ol':
            _process_list(doc, child, images_dir, ordered=True, level=list_level)

        # Tables
        elif tag_name == 'table':
            _process_table(doc, child)

        # Code blocks
        elif tag_name == 'pre':
            code = child.get_text()
            p = doc.add_paragraph(code)
            # Set monospace font for code
            for run in p.runs:
                _set_font(run, is_code=True)

        # Blockquotes
        elif tag_name == 'blockquote':
            _process_element(doc, child, images_dir, list_level)
            # Add some indentation to last paragraph
            if doc.paragraphs:
                doc.paragraphs[-1].paragraph_format.left_indent = Inches(0.5)

        # Horizontal rule
        elif tag_name == 'hr':
            doc.add_paragraph('_' * 50)

        # Divs and other containers - recurse
        elif tag_name in ['div', 'body']:
            _process_element(doc, child, images_dir, list_level)


def _process_inline_elements(paragraph, element, images_dir):
    """
    Process inline elements within a paragraph (bold, italic, links, images, etc.).
    """
    for child in element.children:
        if isinstance(child, str):
            if child.strip() or child == ' ':
                run = paragraph.add_run(child)
                _set_font(run)
            continue

        tag_name = child.name
        text = child.get_text()

        # Strong/Bold
        if tag_name in ['strong', 'b']:
            run = paragraph.add_run(text)
            run.bold = True
            _set_font(run)

        # Emphasis/Italic
        elif tag_name in ['em', 'i']:
            run = paragraph.add_run(text)
            run.italic = True
            _set_font(run)

        # Code (inline)
        elif tag_name == 'code':
            run = paragraph.add_run(text)
            _set_font(run, is_code=True)

        # Links
        elif tag_name == 'a':
            href = child.get('href', '')
            # Add hyperlink (simplified - true hyperlinks in docx require more complex handling)
            run = paragraph.add_run(f'{text} ({href})')
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True
            _set_font(run)

        # Images
        elif tag_name == 'img':
            src = child.get('src', '')
            if images_dir and src:
                _add_image_to_paragraph(paragraph, src, images_dir)

        # Nested inline elements
        else:
            _process_inline_elements(paragraph, child, images_dir)


def _process_list(doc, list_element, images_dir, ordered=False, level=0):
    """
    Process ordered or unordered lists.
    """
    for li in list_element.find_all('li', recursive=False):
        # Get text content, excluding nested lists
        text_parts = []
        for content in li.children:
            if isinstance(content, str):
                text_parts.append(content)
            elif content.name not in ['ul', 'ol']:
                text_parts.append(content.get_text())

        text = ''.join(text_parts).strip()

        if text:
            p = doc.add_paragraph(text, style='List Number' if ordered else 'List Bullet')
            p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
            # Set font for list items
            for run in p.runs:
                _set_font(run)

        # Process nested lists
        for nested in li.find_all(['ul', 'ol'], recursive=False):
            is_ordered = nested.name == 'ol'
            _process_list(doc, nested, images_dir, ordered=is_ordered, level=level + 1)


def _process_table(doc, table_element):
    """
    Convert HTML table to Word table.
    """
    rows = table_element.find_all('tr')
    if not rows:
        return

    # Determine number of columns from first row
    first_row_cells = rows[0].find_all(['th', 'td'])
    num_cols = len(first_row_cells)
    num_rows = len(rows)

    # Create Word table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Fill table
    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for j, cell in enumerate(cells):
            if j < num_cols:  # Safety check
                table.rows[i].cells[j].text = cell.get_text().strip()
                # Set font and make header row bold
                if i == 0 or cell.name == 'th':
                    for paragraph in table.rows[i].cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            _set_font(run)
                else:
                    # Set font for regular cells
                    for paragraph in table.rows[i].cells[j].paragraphs:
                        for run in paragraph.runs:
                            _set_font(run)


def _add_image_to_paragraph(paragraph, image_src, images_dir):
    """
    Add an image to a paragraph, handling relative paths and URL-encoded paths.
    """
    try:
        # URL-decode the image source path (handles Notion exports with encoded Chinese characters)
        decoded_src = unquote(image_src)

        # Handle relative paths
        if not os.path.isabs(decoded_src):
            image_path = os.path.join(images_dir, decoded_src)
        else:
            image_path = decoded_src

        # Check if image exists
        if os.path.exists(image_path):
            # Add image with max width of 6 inches
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(6))
        else:
            # Image not found - add placeholder text
            run = paragraph.add_run(f'[Image not found: {image_src}]')
            run.italic = True
    except Exception as e:
        # Handle any image processing errors
        run = paragraph.add_run(f'[Error loading image: {image_src}]')
        run.italic = True
